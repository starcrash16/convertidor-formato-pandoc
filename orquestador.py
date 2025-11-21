import os
import sys
import subprocess
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# =============================================================================
# 1. CONFIGURACIÓN Y RUTAS
# =============================================================================
BASE = "/data"
ENTRADA_DIR = os.path.join(BASE, "entradas")
ARCHIVOS = os.path.join(BASE, "archivos")
TMP_DIR = os.path.join(BASE, "tmp")

os.makedirs(TMP_DIR, exist_ok=True)

TPL_MASTER = os.path.join(ARCHIVOS, "tpl_master.docx")
LUA_FILTER = os.path.join(ARCHIVOS, "section_break.lua")

VALID_EXT = [".md", ".r", ".R", ".ipynb", ".Rmd", ".tex", ".py"]

# MAPEO DE IMÁGENES
# 0: Portada, 1: Tema, 2: Texto (Hoja 3,5,...), 3: Código (Hoja 4,6,...)
IMG_MAP = {
    0: "header_portada.png",
    1: "header_tema.png",
    2: "header_texto.png",
    3: "header_codigo.png"
}

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
}

# =============================================================================
# 2. MANIPULACIÓN XML (FONDO DETRÁS DEL TEXTO)
# =============================================================================
def make_image_floating_behind_text(run, width_emu, height_emu):
    """Convierte imagen inline a fondo flotante detrás del texto."""
    try:
        inline = run._element.find('.//wp:inline', namespaces=NAMESPACES)
        if inline is None:
            return

        doc_pr = inline.find('.//wp:docPr', namespaces=NAMESPACES)
        if doc_pr is None:
            return
        img_id = doc_pr.get('id')
        img_name = doc_pr.get('name')
        graphic = inline.find('.//a:graphic', namespaces=NAMESPACES)
        if graphic is None:
            return

        anchor_xml = f"""
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" {nsdecls('wp', 'a')}>
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
          <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
          <wp:extent cx="{int(width_emu)}" cy="{int(height_emu)}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          <wp:docPr id="{img_id}" name="{img_name}"/>
          <wp:cNvGraphicFramePr><a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/></wp:cNvGraphicFramePr>
        </wp:anchor>
        """
        anchor_element = parse_xml(anchor_xml)
        anchor_element.append(graphic)
        parent = inline.getparent()
        parent.replace(inline, anchor_element)
    except Exception:
        # No interrumpir el flujo por errores menores en el hack XML
        import traceback; traceback.print_exc()

# =============================================================================
# 3. GESTIÓN DE ENCABEZADOS Y ESTILOS
# =============================================================================
def limpiar_header(header_obj):
    """Elimina párrafos existentes en el encabezado."""
    if header_obj is None:
        return
    for paragraph in list(header_obj.paragraphs):  # usamos list() para iterar de forma segura
        try:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        except Exception:
            pass

def inyectar_imagen_simple(header_obj, img_filename, width, height):
    """Inserta la imagen en el header_obj y aplica el hack XML."""
    if header_obj is None:
        return
    limpiar_header(header_obj)

    p = header_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = 0
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0

    ruta = os.path.join(ARCHIVOS, img_filename)
    if os.path.exists(ruta):
        run = p.add_run()
        # width/height son objetos Length; int(...) devuelve EMU
        run.add_picture(ruta, width=width, height=height)
        make_image_floating_behind_text(run, int(width), int(height))
    else:
        print(f"[WARN] Imagen no encontrada: {img_filename}")

def aplicar_fondo_por_seccion(section, img_name):
    """
    Aplica un único header (no even/odd) a la sección.
    Este enfoque reproduce exactamente la alternancia por SECCIÓN (Hoja 3/4/5..)
    que quieres: cada sección obtiene su propio header independiente.
    """
    # 1. Ajustes base
    section.header_distance = Cm(0)
    section.different_first_page_header_footer = False

    # 2. Desvincular para no arrastrar encabezados previos
    try:
        section.header.is_linked_to_previous = False
    except Exception:
        pass
    try:
        section.footer.is_linked_to_previous = False
    except Exception:
        pass

    # 3. Inyectar imagen en header principal de la sección
    if img_name:
        inyectar_imagen_simple(section.header, img_name, section.page_width, section.page_height)

def procesar_estilos_y_fondos(docx_path):
    print(f"    -> Post-procesando encabezados (Alternancia por sección)...")
    doc = Document(docx_path)
    master = Document(TPL_MASTER)
    master_sections = master.sections

    for i, section in enumerate(doc.sections):
        # Copiar geometría de master
        idx_master = i if i < 2 else 2
        idx_safe = min(idx_master, len(master_sections) - 1)
        m_sect = master_sections[idx_safe]

        section.page_height = m_sect.page_height
        section.page_width = m_sect.page_width
        section.left_margin = m_sect.left_margin
        section.right_margin = m_sect.right_margin
        section.top_margin = m_sect.top_margin
        section.bottom_margin = m_sect.bottom_margin
        section.orientation = m_sect.orientation

        # Decidir imagen según índice de sección (igual que tu código1)
        if i == 0:
            img = IMG_MAP.get(0)
        elif i == 1:
            img = IMG_MAP.get(1)
        else:
            # A partir de la 3era hoja (i=2), alternar por sección:
            # i==2 (Hoja 3) -> texto (IMG_MAP[2])  -> i%2==0 -> texto
            # i==3 (Hoja 4) -> codigo (IMG_MAP[3]) -> i%2==1 -> codigo
            if i % 2 == 0:
                img = IMG_MAP.get(2)
            else:
                img = IMG_MAP.get(3)

        aplicar_fondo_por_seccion(section, img)

    doc.save(docx_path)

# =============================================================================
# 4. CONVERSIÓN Y METADATOS
# =============================================================================
def extraer_metadatos(filepath, comment_char="#"):
    meta = {"titulo": "Sin Título", "autor": "Autor", "fecha": "2025", "tema": "Tema General"}
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            lines = [f.readline() for _ in range(20)]
        for line in lines:
            clean = line.strip()
            if not clean.startswith(comment_char):
                if clean:
                    break
                continue
            content = clean.lstrip(comment_char).strip()
            if content.upper().startswith("TITULO:"):
                meta["titulo"] = content[7:].strip()
            elif content.upper().startswith("AUTOR:"):
                meta["autor"] = content[6:].strip()
            elif content.upper().startswith("FECHA:"):
                meta["fecha"] = content[6:].strip()
            elif content.upper().startswith("TEMA:"):
                meta["tema"] = content[5:].strip()
    except Exception:
        pass
    return meta

def crear_md_final(meta, cuerpo_md):
    return f"""
::: {{custom-style="TituloDocumento"}}
{meta['titulo']}
:::
::: {{custom-style="AutorUnidad"}}
{meta['autor']}
:::
::: {{custom-style="FechaDocumento"}}
{meta['fecha']}
:::
\\newpage
::: {{custom-style="TituloTema"}}
{meta['tema']}
:::
\\newpage
{cuerpo_md}
"""

def convertir_a_md(entrada):
    """
    Convierte entrada a un archivo .md en TMP_DIR y devuelve la ruta al .md generado.
    Para archivos .md se copia el archivo a TMP_DIR y devolvemos la ruta.
    """
    base = os.path.splitext(os.path.basename(entrada))[0]
    ext = os.path.splitext(entrada)[1].lower()
    salida = os.path.join(TMP_DIR, base + ".md")

    cuerpo_md = ""
    meta = {"titulo": base, "autor": "Generado", "fecha": "2025", "tema": "Reporte"}

    if ext == ".md":
        # copiar al tmp y devolver ruta
        with open(entrada, "r", encoding="utf-8") as fin, open(salida, "w", encoding="utf-8") as fout:
            fout.write(fin.read())
        return salida
    elif ext in [".py", ".r"]:
        comment_char = "#"  # para py y r
        meta = extraer_metadatos(entrada, comment_char)
        with open(entrada, "r", encoding="utf-8") as f:
            content = f.read()
        lang = "python" if ext == ".py" else "r"
        cuerpo_md = f"```{lang}\n{content}\n```"
    elif ext == ".tex":
        meta = extraer_metadatos(entrada, "%")
        res = subprocess.run(["pandoc", entrada, "-t", "markdown"], capture_output=True, text=True, check=False)
        cuerpo_md = res.stdout
    elif ext == ".ipynb":
        res = subprocess.run(["pandoc", entrada, "-t", "markdown"], capture_output=True, text=True, check=False)
        cuerpo_md = res.stdout
    else:
        # Otros: intentar convertir con pandoc
        res = subprocess.run(["pandoc", entrada, "-t", "markdown"], capture_output=True, text=True, check=False)
        cuerpo_md = res.stdout

    contenido = crear_md_final(meta, cuerpo_md)
    with open(salida, "w", encoding="utf-8") as f:
        f.write(contenido)
    return salida

# =============================================================================
# 5. PROCESAMIENTO BATCH
# =============================================================================
def obtener_archivos_entrada():
    archivos = []
    if not os.path.exists(ENTRADA_DIR):
        return archivos
    for fname in os.listdir(ENTRADA_DIR):
        if any(fname.lower().endswith(ext.lower()) for ext in VALID_EXT):
            archivos.append(os.path.join(ENTRADA_DIR, fname))
    return archivos

def procesar_archivo(entrada):
    nombre_base = os.path.splitext(os.path.basename(entrada))[0]
    docx_salida = os.path.join(BASE, f"{nombre_base}.docx")
    print(f"\n[INFO] Procesando: {os.path.basename(entrada)}")
    try:
        # obtenemos md_path (si .md lo copiamos; si no lo convertimos)
        md_path = convertir_a_md(entrada) if not entrada.lower().endswith(".md") else os.path.join(TMP_DIR, f"{nombre_base}_temp.md")
        if entrada.lower().endswith(".md"):
            with open(entrada, "r", encoding="utf-8") as f_in, open(md_path, "w", encoding="utf-8") as f_out:
                f_out.write(f_in.read())

        # Generar docx con pandoc
        subprocess.run(["pandoc", md_path, "-o", docx_salida, "--reference-doc", TPL_MASTER, "--standalone", "--lua-filter", LUA_FILTER], check=True)

        # Post-procesar encabezados / fondos por sección
        procesar_estilos_y_fondos(docx_salida)
        print(f"[OK] Generado: {nombre_base}.docx")
        return True
    except subprocess.CalledProcessError as cpe:
        print(f"[ERROR] pandoc falló: {cpe}")
        import traceback; traceback.print_exc()
        return False
    except Exception as e:
        print(f"[ERROR] {e}")
        import traceback; traceback.print_exc()
        return False

def main():
    if not os.path.exists(TPL_MASTER):
        sys.exit(f"[ERROR] Falta {TPL_MASTER}")
    archivos = obtener_archivos_entrada()
    if not archivos:
        print("[WARN] Sin archivos en entradas."); return
    print(f"=== BATCH START: {len(archivos)} archivos ===")
    for f in archivos:
        procesar_archivo(f)
    print("=== BATCH END ===")

if __name__ == "__main__":
    main()

